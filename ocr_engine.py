# ocr_engine.py
import os
import sys
from PIL import Image
import pytesseract
import numpy as np
from pdf2image import convert_from_path
from docx import Document
from docx.shared import RGBColor
from utils import (
    bereinige_zeile, ist_treffer, max_seiten_pro_word_datei,
    preprocess_pillow, preprocess_opencv
)
import shutil

# ---------- Pfad für PyInstaller anpassen ----------
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

# ---------- Tesseract-Pfad ----------
tesseract_cmd = resource_path(os.path.join("Tesseract-OCR", "tesseract.exe"))
pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

# ---------- Poppler-Pfad ----------
poppler_default_path = resource_path(os.path.join("poppler", "Library", "bin"))

# ---------- PDF → PNG ----------
def pdf_to_png(pdf_path, out_dir, poppler_path=None, status_signal=None):
    os.makedirs(out_dir, exist_ok=True)
    if not poppler_path:
        poppler_path = poppler_default_path
    images = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)
    output_files = []
    total_pages = len(images)

    for i, img in enumerate(images):
        out_path = os.path.join(out_dir, f"{os.path.basename(pdf_path)}_{i+1}.png")
        img.save(out_path, "PNG")
        output_files.append(out_path)
        if status_signal:
            status_signal.emit(f"PDF '{os.path.basename(pdf_path)}': Seite {i+1}/{total_pages} konvertiert...")

    return output_files

# ---------- Bild → PNG ----------
def image_to_png(image_path, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(image_path))[0]
    out_path = os.path.join(out_dir, f"{base}.png")
    try:
        img = Image.open(image_path).convert("RGB")
        img.save(out_path, "PNG")
        return out_path
    except Exception as e:
        print(f"[!] Bildkonvertierung fehlgeschlagen für {image_path}: {e}")
        return None

# ---------- Word-Datei speichern ----------
def sichere_datei_speichern(doc, dateiname, output_dir):
    base_name = os.path.splitext(dateiname)[0]
    ext = ".docx"
    full_path = os.path.join(output_dir, dateiname)
    counter = 1
    while True:
        try:
            doc.save(full_path)
            return full_path
        except PermissionError:
            full_path = os.path.join(output_dir, f"{base_name}_{counter}{ext}")
            counter += 1
        except Exception as e:
            print(f"[!] Fehler beim Speichern von {full_path}: {e}")
            return None

# ---------- OCR starten ----------
def starte_ocr(dateien, suchbegriffe, sprache, optimierung=None,
               full_doc=False, highlight=False,
               status_signal=None, progress_signal=None,
               poppler_path=None, output_dir=None,
               abbrechen_flag=None):
    fundstellen = []
    seiten_zaehler = 0
    doc_index = 1
    word_docs = []
    treffer_datei = None
    aktuelles_doc = Document() if full_doc else None
    total_files = len(dateien)

    # temporärer PNG-Ordner für alle Dateien
    png_dir = os.path.join(output_dir, "temp_png") if output_dir else os.path.join(os.getcwd(), "temp_png")
    os.makedirs(png_dir, exist_ok=True)
    temp_files = []

    try:
        for i, datei in enumerate(dateien, start=1):
            if abbrechen_flag and abbrechen_flag():
                if status_signal:
                    status_signal.emit("[!] OCR abgebrochen.")
                return word_docs, treffer_datei

            if status_signal:
                status_signal.emit(f"[{i}/{total_files}] Verarbeite: {os.path.basename(datei)}")
            if progress_signal:
                progress_signal.emit(int((i-1)/total_files*100))

            paths_to_process = []
            if datei.lower().endswith(".pdf"):
                if status_signal:
                    status_signal.emit("   ↳ PDF erkannt, wandle um...")
                pdf_pngs = pdf_to_png(datei, png_dir, poppler_path, status_signal=status_signal)
                paths_to_process.extend(pdf_pngs)
                temp_files.extend(pdf_pngs)
            else:
                png_file = image_to_png(datei, png_dir)
                if png_file:
                    paths_to_process.append(png_file)
                    temp_files.append(png_file)

            for pfad in paths_to_process:
                if abbrechen_flag and abbrechen_flag():
                    if status_signal:
                        status_signal.emit("[!] OCR abgebrochen.")
                    return word_docs, treffer_datei

                if status_signal:
                    status_signal.emit(f"   ↳ OCR: {os.path.basename(pfad)} wird verarbeitet...")

                # -------- Bildoptimierung --------
                try:
                    if optimierung == "pillow":
                        if status_signal:
                            status_signal.emit(f"   ↳ Bildoptimierung mit Pillow...")
                        bild = preprocess_pillow(pfad)
                    elif optimierung == "opencv":
                        if status_signal:
                            status_signal.emit(f"   ↳ Bildoptimierung mit OpenCV...")
                        temp_img = Image.open(pfad).convert("L")
                        temp_np = np.array(temp_img, dtype=np.uint8)
                        if temp_np.size == 0:
                            raise ValueError("Bild ist leer.")
                        bild = preprocess_opencv(temp_np)
                    elif optimierung == "kombiniert":
                        if status_signal:
                            status_signal.emit(f"   ↳ Bildoptimierung kombiniert: Pillow + OpenCV...")
                        temp_img = preprocess_pillow(pfad)
                        temp_np = np.array(temp_img, dtype=np.uint8)
                        if temp_np.size == 0:
                            raise ValueError("Bild ist leer.")
                        bild = preprocess_opencv(temp_np)
                    else:
                        if status_signal:
                            status_signal.emit(f"   ↳ Keine Bildoptimierung...")
                        bild = Image.open(pfad)

                    if np.array(bild).size == 0:
                        raise ValueError("Bild ist leer oder konnte nicht gelesen werden.")

                except Exception as e:
                    fehlertext = f"[!] Bildoptimierung fehlgeschlagen für {pfad}: {e}"
                    if status_signal:
                        status_signal.emit(fehlertext)
                    print(fehlertext)
                    continue  # Bild überspringen

                # OCR starten
                if bild is not None:
                    text = pytesseract.image_to_string(bild, lang=sprache)

                    # -------- Word-Dokument --------
                    if full_doc:
                        if aktuelles_doc.paragraphs:
                            aktuelles_doc.add_page_break()
                        aktuelles_doc.add_heading(os.path.basename(datei), level=1)

                        for line in text.splitlines():
                            if abbrechen_flag and abbrechen_flag():
                                if status_signal:
                                    status_signal.emit("[!] OCR abgebrochen.")
                                return word_docs, treffer_datei

                            p = aktuelles_doc.add_paragraph(line)
                            if highlight:
                                for wort in suchbegriffe:
                                    if wort in line.lower():
                                        start_idx = line.lower().find(wort)
                                        p.clear()
                                        p.add_run(line[:start_idx])
                                        run = p.add_run(line[start_idx:start_idx+len(wort)])
                                        run.font.color.rgb = RGBColor(255, 0, 0)
                                        p.add_run(line[start_idx+len(wort):])

                    # Treffer erfassen
                    for line in text.splitlines():
                        if abbrechen_flag and abbrechen_flag():
                            if status_signal:
                                status_signal.emit("[!] OCR abgebrochen.")
                            return word_docs, treffer_datei

                        treffer = ist_treffer(line, suchbegriffe)
                        if treffer:
                            fundstellen.append(f"{os.path.basename(pfad)}: {', '.join(treffer)} → {bereinige_zeile(line)}")

        # Letztes Word-Dokument speichern
        if full_doc and aktuelles_doc and aktuelles_doc.paragraphs:
            doc_name = f"ocr_ausgabe_{doc_index}.docx"
            gespeicherte_datei = sichere_datei_speichern(aktuelles_doc, doc_name, output_dir)
            if gespeicherte_datei:
                word_docs.append(gespeicherte_datei)

        # Treffer immer als Word-Dokument speichern
        if fundstellen:
            treffer_doc = Document()
            treffer_doc.add_heading("OCR Treffer", level=1)
            for line in fundstellen:
                treffer_doc.add_paragraph(line)
            treffer_name = "ocr_treffer.docx"
            gespeicherte_datei = sichere_datei_speichern(treffer_doc, treffer_name, output_dir)
            if gespeicherte_datei:
                word_docs.append(gespeicherte_datei)
                treffer_datei = gespeicherte_datei

    finally:
        # ---------- temporäre PNG-Dateien löschen ----------
        if temp_files:
            for f in temp_files:
                try:
                    os.remove(f)
                except:
                    pass
            # Optional den Ordner komplett entfernen, wenn leer
            try:
                if os.path.exists(png_dir) and not os.listdir(png_dir):
                    os.rmdir(png_dir)
            except:
                pass

    if status_signal:
        status_signal.emit(f"\n[✔] OCR abgeschlossen. {len(fundstellen)} Treffer.\nErgebnisse: {', '.join(word_docs)}")
    print(f"[✔] OCR abgeschlossen. {len(fundstellen)} Treffer. Ergebnisse: {', '.join(word_docs)}")

    if progress_signal:
        progress_signal.emit(100)

    return word_docs, treffer_datei
